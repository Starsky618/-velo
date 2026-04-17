"""
生成 Starsky 的学习笔记 docx 文档。

用途：
    把 AI 对话中的知识讲解固化成可查阅的 Word 文档，
    作为"我在做 VELO 过程中学到的新知识"的长期积累本。

风格硬要求（v2 修订）：
    语言风格必须严格遵循全局 CLAUDE.md"沟通规则"：说人话、讲小故事、
    用生活类比、术语第一次出现必解释、多用"你""我"对话感。
    不允许用教科书式的抽象 bullet。

使用：
    python3 scripts/gen_learning_notes.py

输出：
    docs/learning/VELO学习笔记.docx

以后怎么加新章节：
    在本脚本的 `add_chapter_N_xxx()` 函数里追加新章节，
    或者直接在生成的 docx 里手动编辑（不再跑脚本覆盖即可）。
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ============================================================
# 字体 & 样式工具（不变）
# ============================================================

CN_FONT = "微软雅黑"
EN_FONT = "Consolas"


def set_font(run, size=11, bold=False, color=None, chinese=True):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = EN_FONT
    if chinese:
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), CN_FONT)


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    sizes = {1: 22, 2: 18, 3: 15, 4: 13}
    set_font(run, size=sizes.get(level, 11), bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return h


def add_para(doc, text, size=11, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = EN_FONT
    run.font.size = Pt(10)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 40)
    set_font(run, size=10, color=RGBColor(0xBB, 0xBB, 0xBB))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10.5, bold=True)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
    if col_widths:
        for row in table.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = Cm(w)
    return table


# ============================================================
# 封面 & 引言（生动版）
# ============================================================


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VELO 学习笔记")
    set_font(run, size=28, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("— Starsky 的成长地图 —")
    set_font(run, size=14, color=RGBColor(0x80, 0x80, 0x80))

    doc.add_paragraph()
    doc.add_paragraph()

    add_heading(doc, "你为什么需要这份笔记？", level=2)

    add_para(
        doc,
        "大学上课学的那些 CS 概念——数据结构、算法、操作系统、数据库——"
        "你可能当时听得明明白白，考完试两周就忘了。为什么？",
    )
    add_para(
        doc,
        "因为这些知识在课堂上是"
        "「抽象的—独立的—没和任何具体东西挂钩的」。"
        "你脑子里没有一个活生生的场景，学完像在空气里画画，风一吹就散。",
        bold=True,
    )
    add_para(
        doc,
        "这份笔记不一样。这里的每一条知识，都和你正在做的 VELO 挂钩："
        "排行榜卡了？那是"
        "「堆」"
        "能解决的事。审查 token 爆炸了？那是"
        "「B+ 树」"
        "的思想。Redis 里的 JWT 缓存？那就是"
        "「哈希表」"
        "在裸奔。",
    )

    doc.add_paragraph()
    add_heading(doc, "学习方法论（三句话说完）", level=2)
    add_bullet(doc, "不要系统刷算法题——你是创业者+大一学生，没那时间")
    add_bullet(
        doc,
        "按"
        "「项目踩坑 → 倒推是哪个知识 → 来这份笔记学 → 回去解决」"
        "的循环走——每个知识点都带着血淋淋的应用场景，印象比课堂深 3 倍",
    )
    add_bullet(
        doc,
        "每学一章，在"
        "「学习动机」"
        "那一行记下「我是因为遇到什么问题才来学这个的」——半年后翻看，每条都是你的"
        "破案档案",
    )

    doc.add_paragraph()
    add_quote(
        doc,
        "一句话心态：别当学生，当侦探。每次项目报错、性能崩、"
        "AI 骂你、审查爆 token，都是一个线索。顺着线索学到的东西，永远比老师塞给你的记得牢。",
    )

    doc.add_page_break()


# ============================================================
# 第 1 章：数据结构与算法速览（生动版）
# ============================================================


def add_chapter_1_data_structures(doc):
    add_heading(doc, "第 1 章：数据结构与算法速览", level=1)

    add_para(
        doc,
        "学习动机：2026-04-17。起因是——我们在讨论"
        "「随着 VELO 项目越写越大，AI 每次审查都要逐行扫全项目，"
        "token 消耗指数级爆炸，能不能学数据库 B+ 树那种索引做分层？」。"
        "讨论过程中发现我很多数据结构概念还是模糊的，"
        "趁这次机会把 9 种常见的、以及和 VELO 相关的，全部打包整理一遍。",
        size=10, color=RGBColor(0x80, 0x80, 0x80),
    )
    doc.add_paragraph()

    # ===== 1.1 =====
    add_heading(doc, "1.1 数据结构是什么？说人话版", level=2)

    add_para(doc, "数据结构说白了就是——「不同形状的盒子」。", bold=True)
    doc.add_paragraph()

    add_para(doc, "你搬家的时候是不是会这样：")
    add_bullet(doc, "衣服 → 塞大纸箱（不怕压，随便扔）")
    add_bullet(doc, "书 → 装小扁箱（重量不至于压垮自己）")
    add_bullet(doc, "碗碟 → 带格子的箱子（防止磕碎）")
    doc.add_paragraph()

    add_para(
        doc,
        "如果你偏要反着来——把 50 本书塞进一个大纸箱，搬起来腰都要断；"
        "把衣服装进带格子的箱子，浪费一半空间。",
    )
    add_para(
        doc,
        "程序里存数据一个道理。你的数据是「用户信息」「通知」「赛段成绩」——"
        "选对盒子，10 万条数据也嗖嗖地跑；"
        "选错盒子，1 万条就能把服务器卡死。",
        bold=True,
    )
    doc.add_paragraph()

    add_heading(doc, "三句话判断选哪种盒子", level=3)
    add_para(
        doc,
        "以下是给你的"
        "「口袋决策树」"
        "——不用记公式，遇到具体问题先拿这三句话套一下：",
    )
    add_bullet(doc, "数据「查得多、改得少」 → 数组 或 哈希表（因为读得快）")
    add_bullet(doc, "数据「改得多、查得少」 → 链表 或 堆（因为改得快）")
    add_bullet(doc, "经常「按顺序或范围查」（比如「本周所有通知」）→ B+ 树 或 跳表")

    doc.add_paragraph()

    # ===== 1.2 =====
    add_heading(doc, "1.2 9 种数据结构一览表", level=2)

    add_para(
        doc,
        "这张表是"
        "「地图」"
        "，不用全记住——先扫一遍知道"
        "「原来还有这些」，后面会挑 3 个你最常遇到的详细讲。",
    )
    doc.add_paragraph()

    headers = ["数据结构", "生活类比（一句话）", "查找速度", "适用场景", "VELO 里用吗"]
    rows = [
        ["数组", "编号整齐的一排储物柜", "O(1) 按号直取", "频繁按位置读", "✅ 轨迹点按顺序存"],
        ["链表", "一串便签，每张写下一张位置", "O(n) 从头找", "频繁插删、不需要跳读", "✅ 通知列表（时间追加）"],
        ["哈希表（HashMap）", "健身房储物柜 + 神奇管理员给你小票", "O(1) 平均", "按 key 精准查", "✅ Redis、小程序本地存储"],
        ["平衡二叉树（AVL）", "严格左右平衡的家族树", "O(log n)", "有序 + 频繁查", "不直接用"],
        ["红黑树", "家族树放松版（允许轻微歪一点）", "O(log n)", "高频插删的有序集合", "间接（Linux 内核在用）"],
        ["B 树", "每个房间多个抽屉的家族树", "O(log n)", "减少磁盘访问次数", "间接（数据库索引）"],
        ["B+ 树", "B 树升级版：叶子存货 + 叶子互相手拉手", "O(log n) + 范围快", "数据库索引、范围查询", "✅ PostgreSQL 所有索引都是它"],
        ["堆 / 跳表", "金字塔顶尖永远最大 / 多层目录的链表", "O(log n)", "TopK 查询、优先队列", "✅ Redis ZSet（未来排行榜用）"],
        ["Trie 字典树", "按前缀分叉的树（像通讯录翻页）", "O(k) k=字符长度", "前缀匹配、自动补全", "❌ 未用（赛段搜索可以用）"],
    ]
    add_table(doc, headers, rows, col_widths=[3.0, 4.2, 2.4, 3.2, 3.5])

    doc.add_paragraph()

    # ===== 1.3 =====
    add_heading(doc, "1.3 三个重点深讲：你最常遇到的", level=2)

    # ---------- 哈希表 ----------
    add_heading(doc, "🔑 哈希表 —— 健身房储物柜的神奇管理员", level=3)

    add_heading(doc, "先讲个场景", level=4)
    add_para(
        doc,
        "你去健身房存包。你把包往储物柜那边一放，"
        "管理员不用看、不用找、不用翻本子，"
        "直接撕一张小票塞你手里："
        "「317 号柜」"
        "。",
    )
    add_para(
        doc,
        "一小时后你来取包。你拿着小票"
        "「317 号」"
        "径直走到那个柜子——"
        "连 0.1 秒都不用浪费找它在哪。",
    )
    add_para(
        doc,
        "现在问题来了：这个健身房有 10 万个柜子。管理员怎么做到"
        "「瞬间」"
        "给你分柜号的？他不可能记住"
        "「今天早上谁放了什么」"
        "——10 万个他哪记得过来。",
    )
    add_para(
        doc,
        "他用的是一个"
        "「公式」"
        "——你的会员卡号（比如 K2025-18739）经过一个固定的数学运算，"
        "算出一个 0 到 99999 之间的数字，那个数字就是你的柜号。",
        bold=True,
    )
    add_para(
        doc,
        "这个"
        "「把任何输入算成一个固定范围数字」"
        "的魔法公式，就叫"
        "「哈希函数」"
        "（hash function）。"
        "你可以先不用管它具体怎么算，只要知道它是个确定性的"
        "「输入→输出」"
        "转换器就行。",
    )

    add_heading(doc, "所以哈希表到底是什么？", level=4)
    add_para(doc, "把上面的故事翻译成程序员能懂的话就是：")
    add_bullet(doc, "柜子 = 一排固定位置的储物空间（一个数组）")
    add_bullet(doc, "小票 = 你指定的 key（比如「用户42的JWT」）")
    add_bullet(doc, "哈希函数 = 把 key 算成一个数组位置")
    add_bullet(doc, "存数据 = 柜子[位置] = 你的数据")
    add_bullet(doc, "取数据 = 柜子[位置] → 拿回数据")

    add_heading(doc, "为啥它快到飞起？", level=4)
    add_para(
        doc,
        "因为找东西永远只要一步——算一次哈希函数，直接到柜子门口。"
        "不管你存了 100 条数据，还是 1 亿条数据，"
        "查找速度一模一样。",
        bold=True,
    )
    add_para(
        doc,
        "这就是"
        "「时间复杂度 O(1)」"
        "的意思——「和数据量无关，永远一步到位」。",
    )

    add_heading(doc, "代价是什么？（每个好东西都有代价）", level=4)
    add_para(
        doc,
        "哈希表的代价叫"
        "「哈希冲突」"
        "：有时候管理员的公式会"
        "「算错」"
        "——两个不同的会员算出同一个柜号。",
    )
    add_para(
        doc,
        "两个人都说"
        "「我包在 317 号柜」"
        "，柜子就一个，怎么办？最常见的解决办法：在 317 号柜里挂一条小链子，"
        "两个人的包一前一后挂着，"
        "找的时候到 317 号柜再沿链子挨个看。",
    )
    add_para(
        doc,
        "这时候查找速度就从"
        "「一步到位」"
        "退化成"
        "「一个个翻」"
        "——从 O(1) 退化成 O(n)。"
        "好的哈希函数能让冲突极少发生，所以实际上哈希表平均还是 O(1)。",
    )

    add_heading(doc, "VELO 里它在哪", level=4)
    add_bullet(
        doc,
        "Redis 缓存 = 一张巨大的哈希表。"
        "你绑 Strava 时我们存的 strava_state:{nonce}——"
        "就算 Redis 里有 1000 万个 state，查起来还是一眼的事。",
    )
    add_bullet(
        doc,
        "小程序本地存储 wx.setStorageSync('mute_notifications', true) "
        "——key 是字符串，value 是布尔值，就是哈希表。",
    )
    add_bullet(
        doc,
        "Python 的 dict、JavaScript 的 Object、Java 的 HashMap"
        "——全都是哈希表的不同包装。"
        "你写 dict['key'] 这一行代码，"
        "背后就是"
        "「算哈希 → 去柜号 → 拿数据」"
        "这套操作。",
    )

    add_quote(
        doc,
        "记住这一张图就够了：key → 哈希函数 → 柜号 → 直接取。"
        "后面的一切哈希表变种、碰撞解决、负载因子……都只是这张图的装修。",
    )

    doc.add_paragraph()

    # ---------- B+ 树 ----------
    add_heading(doc, "🌳 B+ 树 —— 数据库躲在背后的秘密武器", level=3)

    add_heading(doc, "先讲个场景", level=4)
    add_para(
        doc,
        "你去图书馆找"
        "《阿Q正传》"
        "。图书馆有 50 万本书。你会怎么找？",
    )
    add_para(
        doc,
        "肯定不会从第一本翻到最后一本（那是 O(n)，鬼知道要多久）。"
        "你直接去检索台——那里有个大柜子，每个小抽屉贴着"
        "「A-B」「C-D」「E-F」"
        "这样的标签。",
    )
    add_para(
        doc,
        "你拉开"
        "「L-M」"
        "（鲁迅的鲁，L 开头）那一抽屉。抽屉里又有很多小卡片，"
        "卡片上贴着"
        "「Lu 鲁」「Li 李」「Liu 刘」"
        "——你拉开"
        "「Lu」"
        "这一张，看到"
        "「鲁迅 - 阿Q正传 - 位置 12-B-347」"
        "。",
    )
    add_para(
        doc,
        "整个过程你翻了 3-4 层（大柜子 → 抽屉 → 卡片 → 书）。"
        "50 万本书，你最多 20 步就找到了——这就是 B+ 树的精髓。",
        bold=True,
    )

    add_heading(doc, "B+ 树到底长什么样？", level=4)
    add_para(doc, "它是一棵"
        "「多叉树」"
        "，但结构比普通树有两个特殊的讲究：")
    add_bullet(
        doc,
        "每个节点（就是每层的盒子）里装着多个 key"
        "——不是二叉树那种一个节点只一个值。"
        "这样树的高度大幅压缩（50 万数据只要 3-4 层）。",
    )
    add_bullet(
        doc,
        "只有最底层的叶子节点才真的存数据，"
        "上层都只是"
        "「目录指针」"
        "。而且——这是 B+ 树最骚的设计——叶子节点彼此之间用链表手拉手连着。",
    )

    add_heading(doc, "为啥这两个设计这么聪明？", level=4)

    add_para(doc, "设计 1：多叉（不是二叉）", bold=True)
    add_para(
        doc,
        "数据库的数据放在硬盘上，而硬盘读一次数据就叫"
        "「磁盘 I/O」"
        "——这玩意儿慢得离谱，大概比内存慢 10 万倍。"
        "所以减少磁盘 I/O 次数比减少"
        "「步数」"
        "重要得多。",
    )
    add_para(
        doc,
        "B+ 树一次磁盘 I/O 可以把一个节点（含几百个 key）整个读进内存，"
        "然后在内存里对比。二叉树一个节点只一个 key，要多读几十倍磁盘——慢 10 万倍的东西多读 10 倍就是 100 万倍慢。"
        "所以 B+ 树比二叉树在数据库场景快太多了。",
    )

    add_para(doc, "设计 2：叶子手拉手", bold=True)
    add_para(
        doc,
        "你在 VELO 里查"
        "「这周所有通知」"
        "——数据库会先用树结构定位到"
        "「这周第一条」"
        "，然后沿着叶子链表一路扫到"
        "「这周最后一条」"
        "。",
    )
    add_para(
        doc,
        "注意：它不用再从树根往下查第二次、第三次——一条"
        "「横向链表」"
        "解决所有范围查询。这就是为什么数据库做"
        "WHERE created_at BETWEEN A AND B"
        "那么快。",
    )

    add_heading(doc, "VELO 里它在哪", level=4)
    add_bullet(
        doc,
        "PostgreSQL 里你建的所有索引，默认都是 B+ 树。"
        "比如我们第 4 期加的 idx_notifications_user_unread"
        "——底层就是 B+ 树。",
    )
    add_bullet(
        doc,
        "查"
        "「小明的所有未读通知」"
        "，数据库做的事：先用树结构定位到小明的第一条未读"
        "→ 沿着叶子链表扫到最后一条未读。几毫秒搞定。",
    )
    add_bullet(
        doc,
        "我们讨论的"
        "「代码审查索引」"
        "——给代码符号（函数名、字段名）建一份 SYMBOLS.md，"
        "就是用了 B+ 树"
        "「分层 + 叶子手拉手」"
        "的思想，"
        "让 AI 不用每次全扫代码。",
    )

    add_quote(
        doc,
        "B+ 树的精髓就是"
        "「冷热分离 + 分层下钻」"
        "：根和中间层永远在内存（像地图），叶子才是真数据（像地图标记的具体地址）。"
        "你写 SQL 时这一切都藏在数据库里——现在你知道它在你背后默默做了什么。",
    )

    doc.add_paragraph()

    # ---------- 堆 + 跳表 ----------
    add_heading(doc, "⛰️ 堆 + 跳表 —— 排行榜背后的两种武器", level=3)

    add_heading(doc, "先讲个场景", level=4)
    add_para(
        doc,
        "你开的骑行 App VELO 有 10 万用户。每天都有人刷出新的赛段成绩。"
        "你想实时维护一个"
        "「青云路赛段 TOP20」"
        "榜单——"
        "新成绩一进来就更新，用户点开看永远是最新前 20。",
    )
    add_para(doc, "你会怎么做？")
    add_para(doc, "有三种方式：", bold=True)
    add_bullet(doc, "方式 A（最蠢）：每次有新成绩就把 10 万条全排序，取前 20。每次 O(n log n)，服务器哭了。")
    add_bullet(doc, "方式 B（中等）：把所有成绩放数据库，每次 ORDER BY + LIMIT 20。慢但勉强能用（VELO 现在就这么做）。")
    add_bullet(doc, "方式 C（高级）：维护一个「只装前 20 名」的结构，新成绩来了只比较它和当前第 20 名。这就要用到「堆」。")

    add_heading(doc, "堆（Heap）—— 金字塔形的容器", level=4)
    add_para(
        doc,
        "堆是一种长得像金字塔的树结构。"
        "它保证一件事：塔顶永远是"
        "「最大（或最小）的那个」"
        "。",
    )
    add_para(doc, "插入和删除塔顶都是 O(log n)——很快。")
    add_para(
        doc,
        "TopK 问题（找前 K 大）怎么用堆解决？"
        "维护一个"
        "「大小固定为 K」"
        "的小顶堆（塔顶永远是这 K 个里最小的）。"
        "每来一条新成绩：",
    )
    add_bullet(doc, "比塔顶（第 K 名）大 → 踢掉塔顶，把新值放进堆 → O(log K)")
    add_bullet(doc, "比塔顶小 → 直接扔掉不管")
    add_para(
        doc,
        "这样无论有多少成绩来，堆里永远只留前 K 名，内存占用固定，处理速度恒定。",
        bold=True,
    )

    add_heading(doc, "跳表（Skip List）—— 地铁快慢车原理", level=4)
    add_para(
        doc,
        "跳表长得像一个"
        "「多层的链表」"
        "——很像地铁有快车和慢车。",
    )
    add_para(doc, "你想象一下北京地铁 4 号线：")
    add_bullet(doc, "最底层：每一站都停（慢车）")
    add_bullet(doc, "中间层：只停每 4 站（快车）")
    add_bullet(doc, "最顶层：只停每 16 站（特快车）")

    add_para(
        doc,
        "从「北京南站」去「海淀黄庄」——你先坐特快到大方向，再换快车再换慢车。"
        "总站数是 O(log n)，比一站站坐（O(n)）快得多。",
    )
    add_para(
        doc,
        "跳表就是这个思路：多层链表，每层跳着连，查找/插入/删除都是 O(log n)。"
        "它和红黑树实现的效果类似，但代码简单得多——"
        "所以 Redis 的 ZSet（有序集合）就用跳表做底层。",
    )

    add_heading(doc, "VELO 里它们在哪", level=4)
    add_bullet(
        doc,
        "现在没用——排行榜走 SQL（方式 B），性能还够。",
    )
    add_bullet(
        doc,
        "未来要用——等用户到百万级、赛段上千条，"
        "SQL 就慢了。到时候把"
        "「每个赛段 TOP20」"
        "放 Redis ZSet，就在用跳表。",
    )
    add_bullet(
        doc,
        "举个更常见的例子——微博热搜榜、抖音评论「热评」、游戏战力榜，"
        "底层基本都是 Redis ZSet = 跳表。",
    )

    add_quote(
        doc,
        "堆适合"
        "「只关心前 K 个」"
        "的场景（塔顶够用）；"
        "跳表适合"
        "「要按顺序查+经常插删」"
        "的场景（快慢车任你跳）。"
        "两个结构在排行榜场景经常一起出现。",
    )

    doc.add_paragraph()

    # ===== 1.4 =====
    add_heading(doc, "1.4 O(log n) 直觉训练 —— 为啥这三个字母这么神", level=2)

    add_heading(doc, "先玩个小游戏", level=3)
    add_para(
        doc,
        "我心里想一个 1 到 1000 之间的整数，你来猜。"
        "你每猜一次，我说"
        "「大了」"
        "或"
        "「小了」"
        "，你再猜。",
    )
    add_para(doc, "玩法 A：从 1 开始一个个猜——「1？2？3？……」")
    add_para(doc, "最坏情况猜 1000 次。这叫 O(n)——线性时间。随着 n 变大，步数成正比变多。")

    add_para(doc, "玩法 B：聪明点——每次猜范围的中间值", bold=True)
    add_bullet(doc, "第 1 次：500（砍一半）")
    add_bullet(doc, "大了 → 第 2 次：250（再砍一半）")
    add_bullet(doc, "小了 → 第 3 次：375（再砍一半）")
    add_bullet(doc, "……")
    add_para(
        doc,
        "最多猜 10 次就一定找到——"
        "因为 2 的 10 次方 = 1024。"
        "这就叫 O(log n)——对数时间。",
        bold=True,
    )

    add_heading(doc, "为啥差这么多", level=3)
    headers = ["数据规模 n", "O(n) 要多少步", "O(log n) 要多少步"]
    rows = [
        ["100", "100", "7"],
        ["1 万", "1 万", "14"],
        ["1 亿", "1 亿", "约 30"],
        ["1 万亿（世界所有数据）", "1 万亿", "约 40"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 4, 5])

    add_para(
        doc,
        "看见没——1 亿 vs 1 万亿，O(n) 是 1 万倍差距，"
        "但 O(log n) 只多 10 步。"
        "这就是为什么所有"
        "「能做到 O(log n)」"
        "的数据结构（B+ 树、跳表、堆、平衡二叉树）都是工程师的宝贝。",
        bold=True,
    )

    add_quote(
        doc,
        "记住一句话：「n 翻倍，步数只加 1」——这就是 log 的魔法。"
        "当你听说一个算法是 O(log n)，默认它是够快的。",
    )

    doc.add_paragraph()

    # ===== 1.5 =====
    add_heading(doc, "1.5 VELO 里每个地方都在用什么结构", level=2)

    add_para(
        doc,
        "下面这张表是"
        "「你打开 VELO 写代码时，不知不觉就在使用的数据结构」"
        "——看完你会发现，原来这些知识每天都在帮你干活。",
    )
    doc.add_paragraph()

    headers = ["你在做的事", "背后用了什么结构", "目前状态"]
    rows = [
        ["AI 代码审查想减少全扫 token", "B+ 树思想（做符号表做中间层）", "🚧 规划中"],
        ["查「某个字段在代码里存不存在」", "哈希表（符号表本质就是 HashMap）", "🚧 规划中"],
        ["检查模块之间有没有循环引用", "图 + DFS 深度优先遍历", "🚧 可做"],
        ["用户输入赛段名做前缀搜索", "Trie 字典树", "❌ 未用"],
        ["通知列表按时间倒序显示", "链表（只往头追加）", "✅ 用了"],
        ["赛段 TOP20 排行榜", "堆 / 跳表（目前 SQL 顶着）", "⚠️ 百万用户再优化"],
        ["JWT / OAuth state 缓存", "哈希表（Redis 就是它）", "✅ 用了"],
        ["数据库里的所有索引", "B+ 树（PostgreSQL 默认）", "✅ 用了，躲在背后"],
    ]
    add_table(doc, headers, rows, col_widths=[6.5, 6, 3])

    add_quote(
        doc,
        "意识到这一点很重要：你已经每天都在用这些数据结构了——"
        "只是没意识到而已。学数据结构不是"
        "「学新东西」"
        "，而是"
        "「搞清楚你一直在用的东西到底是什么」"
        "。",
    )

    doc.add_paragraph()

    # ===== 1.6 =====
    add_heading(doc, "1.6 我的学习路径建议（三选一，我强推第一条）", level=2)

    add_heading(doc, "🥇 路 1（强推）：带着项目问题学", level=3)
    add_para(
        doc,
        "每次 VELO 踩坑 → 先停下来想"
        "「这个坑是哪个数据结构不合适？」"
        "→ 去学那一种 → 用上 → 写进这份笔记。",
    )
    add_para(doc, "举几个你未来会遇到的例子：", bold=True)
    add_bullet(doc, "排行榜慢到 10 秒才出结果 → 意识到「原来我该用堆」 → 学堆 → 用 Redis ZSet 重构")
    add_bullet(doc, "审查 token 每次都爆 → 意识到「这是 O(n^2) 的问题」 → 学 B+ 树 → 做符号表索引")
    add_bullet(doc, "赛段搜索反应慢 → 意识到「前缀匹配该用 Trie」 → 学 Trie → 加进搜索模块")
    add_para(
        doc,
        "这种学法有个大优势——每个知识点都带着血淋淋的"
        "「我为什么要学这个」"
        "的记忆。印象比课堂深 3 倍。",
        bold=True,
    )

    add_heading(doc, "🥈 路 2：系统性补基础（周末用）", level=3)
    add_bullet(doc, "《算法图解》——入门书，图多，2-3 天就能读完，入门的最好选择")
    add_bullet(doc, "《算法导论》——圣经级工具书，别从头读，当查阅字典用")
    add_bullet(doc, "LeetCode Easy 前 100 题——实战，但不要沉迷刷题，刷 30 题有感觉就停")

    add_heading(doc, "🥉 路 3：读工程源码（高阶）", level=3)
    add_bullet(doc, "PostgreSQL 索引源码 → 看 B+ 树真实怎么实现的")
    add_bullet(doc, "Redis 数据结构源码 → 字符串/列表/哈希/集合/ZSet 的真实底层")
    add_bullet(doc, "Linux 内核红黑树实现 → 调度器里怎么用红黑树管理进程")

    doc.add_paragraph()
    add_quote(
        doc,
        "给你的建议：按路 1 走。你现在是创业者+大一学生+ vibe coder，"
        "没时间系统刷题。每次项目踩坑就来这份笔记加一章，"
        "半年 = 6 章，一年 = 12 章——"
        "到时候你的数据结构地图会比很多大三学生完整。",
    )

    add_divider(doc)


# ============================================================
# 附：未来章节占位
# ============================================================


def add_future_chapters(doc):
    add_heading(doc, "附：未来章节占位（遇到新坑就来这里加）", level=1)

    add_para(
        doc,
        "每一章按下面这个结构组织，保持和第 1 章一样的"
        "「讲故事 + 生活类比 + VELO 挂钩 + 路径建议」"
        "风格。不要写成教科书。",
        bold=True,
    )

    add_heading(doc, "模板结构", level=2)
    add_bullet(doc, "X.1 为什么要学这个？ —— 起源的项目痛点（讲个具体场景）")
    add_bullet(doc, "X.2 这是什么？说人话版 —— 生活类比 + 通俗定义")
    add_bullet(doc, "X.3 核心对比大表 / 概念速览 —— 地图式全景")
    add_bullet(doc, "X.4 重点深讲 —— 挑 2-3 个最相关的，每个讲「先讲场景 + 原理 + 代价 + VELO 在哪用」")
    add_bullet(doc, "X.5 直觉训练 —— 数字对比 / 生活类比让概念落地")
    add_bullet(doc, "X.6 VELO 工作流映射 —— 让你看到「原来我每天都在用」")
    add_bullet(doc, "X.7 学习路径建议 —— 按「项目倒推」原则推路径")

    doc.add_paragraph()

    add_heading(doc, "下一章候选主题（按优先级）", level=2)
    add_para(
        doc,
        "这些是你大概率会遇到、值得单开一章的主题。按优先级排——"
        "越靠前越可能先撞上。",
    )
    add_bullet(
        doc,
        "操作系统进程/线程/协程 —— VELO 的 Worker 和 Scheduler 就是进程，"
        "以后异步 I/O 会遇到协程",
    )
    add_bullet(
        doc,
        "数据库事务/隔离级别/锁 —— 你已经踩过 SAVEPOINT 的坑，"
        "下次遇到并发问题会需要理解"
        "「脏读 / 幻读 / 不可重复读」",
    )
    add_bullet(
        doc,
        "网络 HTTP/HTTPS/WebSocket/DNS —— Strava OAuth 跨域、"
        "Webhook 来源校验都是网络层的事",
    )
    add_bullet(
        doc,
        "设计模式 观察者/工厂/策略 —— notification 的 detect_events 就是典型"
        "观察者模式",
    )
    add_bullet(
        doc,
        "分布式 CAP/一致性哈希/限流/熔断 —— 用户量上来后必学，现在先不急",
    )

    doc.add_paragraph()
    add_quote(
        doc,
        "别急着把这些全学完。"
        "每次项目遇到新坑，就认领一个主题来这份笔记里加一章。"
        "半年 6 章，一年 12 章——"
        "这就是属于你自己的、带着血汗味的工程师手册。",
    )


# ============================================================
# 主流程
# ============================================================


def build_document(out_path: Path):
    doc = Document()

    # 全局默认字体
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)

    add_cover(doc)
    add_chapter_1_data_structures(doc)
    doc.add_page_break()
    add_future_chapters(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"✅ 生成完毕：{out_path}")


if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "docs" / "learning" / "VELO学习笔记.docx"
    build_document(out)
